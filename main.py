import os
import re
from tkinter import Tk, Label, Entry, Button, Listbox, filedialog, messagebox, MULTIPLE
from tkinter.ttk import Combobox
from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Define the DocumentUpdater class
class DocumentUpdater:
    def __init__(self, file_path):
        self.doc = Document(file_path)
        self.file_path = file_path

    # Method to replace '$@$' with a new word throughout the document
    def replace_word(self, old_word, new_word):
        print(f"Replacing all occurrences of '{old_word}' with '{new_word}'...")
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if old_word in run.text:
                    run.text = run.text.replace(old_word, new_word)
                    run.bold = True  # Apply bold formatting to the new word
                    print(f"Replaced in paragraph: '{paragraph.text}'")

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_word in cell.text:
                        cell.text = cell.text.replace(old_word, new_word)
                        print(f"Replaced in table cell: '{cell.text}'")

    # Method to update the footer reference and align text to the right
    def update_footer_reference(self, old_reference, new_reference):
        print(f"Updating reference from '{old_reference}' to '{new_reference}'...")
        reference_updated = False
        for section in self.doc.sections:
            footer = section.footer
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if old_reference in cell.text:
                            print(f"Original reference in footer: '{cell.text}'")
                            cell.text = new_reference  # Overwrite the entire reference to avoid duplication
                            reference_updated = True
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align the text to the right
                            print(f"Updated reference in footer to: '{cell.text}'")
                            break  # Exit the loop after updating

        if not reference_updated:
            print(f"No reference found with '{old_reference}' to replace.")

    # Method to update the date in the first-page table where the second row contains the date
    def update_date_in_table(self):
        current_date = datetime.now().strftime("%b %d %Y")
        print(f"Updating the date to '{current_date}'...")
        date_updated = False
        for table in self.doc.tables:
            for row_index, row in enumerate(table.rows):
                if "Submission Date" in row.cells[0].text:
                    if row_index + 1 < len(table.rows):
                        date_cell = table.rows[row_index + 1].cells[0]
                        print(f"Original date found: '{date_cell.text}'")
                        date_cell.text = ""
                        run = date_cell.paragraphs[0].add_run(current_date)
                        run.font.size = Pt(9)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        date_updated = True
                        print(f"Date updated to: '{current_date}'")
                        break

        if not date_updated:
            print(f"No date found to update.")

    # Method to update the term "YEAR" in the terms and conditions, making only the quoted part bold
    def update_contractual_year(self, new_year):
        for paragraph in self.doc.paragraphs:
            if "Minimum Commitment Contractual Period" in paragraph.text:
                print(f"Updating contractual year in paragraph: '{paragraph.text}'")
                
                # Find the text inside the quotes using regex
                match = re.search(r'“[^”]+”', paragraph.text)  # Match the text inside the double quotation marks
                if match:
                    start, end = match.span()  # Get the start and end index of the match
                    before_quotes = paragraph.text[:start]
                    inside_quotes = f'“{new_year} YEAR”'  # The new bold text
                    after_quotes = paragraph.text[end:]

                    # Clear the paragraph and reconstruct it with formatting
                    paragraph.clear()
                    
                    # Add the part before the quotes (not bold)
                    run_before = paragraph.add_run(before_quotes)
                    run_before.font.size = Pt(9)  # Ensure it's not bold

                    # Add the part inside the quotes (bold)
                    run_inside = paragraph.add_run(inside_quotes)
                    run_inside.font.size = Pt(9)
                    run_inside.bold = True  # Make the text inside quotes bold

                    # Add the part after the quotes (not bold)
                    run_after = paragraph.add_run(after_quotes)
                    run_after.font.size = Pt(9)

                    print(f"Updated contractual period to: '{paragraph.text}'")
                break

    # Method to save the updated document with a sanitized filename
    def save_document(self, new_reference, save_location):
        file_name = new_reference.replace('Ref.: ', '')  # Clean the file name
        file_name = re.sub(r'\s+', '', file_name)
        file_name = re.sub(r'[^\w\s]', '_', file_name)
        file_name = re.sub(r'_+', '_', file_name)
        file_name = file_name.strip('_')

        # Save the document in the specified location
        file_path = os.path.join(save_location, f'{file_name}.docx')
        if os.path.exists(file_path):
            os.remove(file_path)

        self.doc.save(file_path)
        print(f"Document saved successfully as '{file_path}'!")
        return file_path

# GUI Application using Tkinter
class DocumentProcessorApp:
    def __init__(self, master):
        self.master = master
        master.title("Document Processor GUI")

        # Labels and Textboxes
        Label(master, text="New Company Name:").grid(row=0, column=0)
        self.company_name_entry = Entry(master)
        self.company_name_entry.grid(row=0, column=1)

        Label(master, text="Reference Number:").grid(row=1, column=0)
        self.reference_number_entry = Entry(master)
        self.reference_number_entry.grid(row=1, column=1)

        Label(master, text="Add new word:").grid(row=2, column=0)
        self.new_word_entry = Entry(master)
        self.new_word_entry.grid(row=2, column=1)

        Label(master, text="Select Product Name(s):").grid(row=3, column=0)
        self.product_listbox = Listbox(master, selectmode=MULTIPLE, height=4)
        self.product_listbox.grid(row=3, column=1)
        self.product_listbox.insert("end", "Industrial Router", "Mdawm", "MDM")

        Label(master, text="Contractual Year:").grid(row=4, column=0)
        self.contractual_year_entry = Entry(master)
        self.contractual_year_entry.grid(row=4, column=1)

        # Drop-down for template selection
        Label(master, text="Select Template File:").grid(row=5, column=0)
        self.template_combobox = Combobox(master)
        self.template_combobox.grid(row=5, column=1)

        Button(master, text="Load Templates", command=self.load_templates).grid(row=6, column=1)

        # Buttons
        Button(master, text="Process Document", command=self.process_document).grid(row=7, column=0, columnspan=2)

        self.output_label = Label(master, text="")
        self.output_label.grid(row=8, column=0, columnspan=2)

        self.file_paths = {}  # Dictionary to store full file paths with filenames
        self.selected_product_indices = []  # Store product selections

    def load_templates(self):
        # **Important Fix**: Store the current product selections
        self.selected_product_indices = self.product_listbox.curselection()

        # Load the templates into the template combobox
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.file_paths.clear()  # Clear the previous file paths
            template_names = []
            for file in os.listdir(folder_path):
                if file.endswith(".docx"):
                    template_names.append(file)
                    self.file_paths[file] = os.path.join(folder_path, file)  # Store full path for processing
            self.template_combobox['values'] = template_names  # Add template names to the combobox

        # **Important Fix**: Restore product selections after loading templates
        for i in self.selected_product_indices:
            self.product_listbox.selection_set(i)

    def process_document(self):
        selected_template = self.template_combobox.get()
        if selected_template:
            try:
                file_path = self.file_paths[selected_template]  # Get full file path from dictionary
                new_word = self.new_word_entry.get() or '$@$'  # Use $@$ if no new word is entered
                company_name = self.company_name_entry.get()
                reference_number = self.reference_number_entry.get()
                contractual_year = self.contractual_year_entry.get()

                # Process selected products and join them with '&'
                selected_products = [self.product_listbox.get(i) for i in self.product_listbox.curselection()]
                product_name = "&".join(selected_products)

                # Generate new reference
                new_reference = f'Ref.: {company_name}_{product_name}_{reference_number}_{datetime.now().strftime("%d%m%y")}'

                # Save location set in the code (instead of asking every time)
                save_location = r'C:\Users\ahmed\Desktop\presales files\app\templates\save'

                updater = DocumentUpdater(file_path)
                updater.replace_word('$@$', new_word)  # Replace '$@$' with the new word
                updater.update_footer_reference('Ref.: ', new_reference)  # Update the footer
                updater.update_date_in_table()  # Update date in the first-page table
                if contractual_year:
                    updater.update_contractual_year(contractual_year)  # Update contractual year
                updater.save_document(new_reference, save_location)  # Save the document with the new reference
                messagebox.showinfo("Success", f"Document processed and saved successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred: {str(e)}")
        else:
            messagebox.showerror("Error", "Please select a template to process.")

if __name__ == "__main__":
    root = Tk()
    app = DocumentProcessorApp(root)
    root.mainloop()
